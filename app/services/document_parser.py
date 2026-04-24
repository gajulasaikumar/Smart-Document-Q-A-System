from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass(slots=True)
class ExtractedBlock:
    text: str
    page_number: int | None = None


def parse_document(file_path: Path, extension: str) -> list[ExtractedBlock]:
    if extension == ".pdf":
        return _parse_pdf(file_path)
    if extension == ".docx":
        return _parse_docx(file_path)
    raise ValueError(f"Unsupported document extension: {extension}")


def _clean_text(value: str) -> str:
    return " ".join(value.split()).strip()


def _parse_pdf(file_path: Path) -> list[ExtractedBlock]:
    reader = PdfReader(str(file_path))
    blocks: list[ExtractedBlock] = []

    for page_number, page in enumerate(reader.pages, start=1):
        text = _clean_text(page.extract_text() or "")
        if text:
            blocks.append(ExtractedBlock(text=text, page_number=page_number))

    if not blocks:
        raise ValueError("No text could be extracted from the PDF document.")
    return blocks


def _parse_docx(file_path: Path) -> list[ExtractedBlock]:
    document = DocxDocument(str(file_path))
    blocks: list[ExtractedBlock] = []

    for paragraph in document.paragraphs:
        text = _clean_text(paragraph.text)
        if text:
            blocks.append(ExtractedBlock(text=text))

    for table in document.tables:
        for row in table.rows:
            cells = [_clean_text(cell.text) for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                blocks.append(ExtractedBlock(text=row_text))

    if not blocks:
        raise ValueError("No text could be extracted from the DOCX document.")
    return blocks
